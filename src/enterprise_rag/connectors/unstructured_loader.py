from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument

from enterprise_rag.models import Document


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 120) -> list[str]:
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end].strip())
        start = end - overlap
        if start < 0:
            start = 0
        if end == len(text):
            break
    return [c for c in chunks if c]


def _read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in [".txt", ".md"]:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    return ""


def load_unstructured_documents(data_dir: Path) -> list[Document]:
    documents: list[Document] = []
    supported = {".txt", ".md", ".pdf", ".docx"}

    for path in sorted(data_dir.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in supported:
            continue

        text = _read_file(path).strip()
        if not text:
            continue

        chunks = _chunk_text(text)
        for idx, chunk in enumerate(chunks):
            documents.append(
                Document(
                    id=f"{path.stem}-{idx}",
                    text=chunk,
                    metadata={
                        "source_type": "unstructured",
                        "path": str(path),
                        "chunk_id": idx,
                    },
                )
            )
    return documents

